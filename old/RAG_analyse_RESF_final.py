import os
import time
import torch
import httpx
from transformers import AutoTokenizer, AutoModel
from huggingface_hub import login
from langchain.schema import Document as LangDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_mistralai import MistralAIEmbeddings, ChatMistralAI
from docx import Document as WordDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================ Clés API ============================
os.environ["HF_TOKEN"] = "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"
key_Mistral = "ThSJGAwoQibJjTiNJq1YjLdVyGvtmpeM"
login(os.environ["HF_TOKEN"])

# ============================ Modèle d'embeddings local ============================
model = AutoModel.from_pretrained("BAAI/bge-m3")
tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-m3")

def get_embeddings(texts):
    inputs = tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
    with torch.no_grad():
        embeddings = model(**inputs).last_hidden_state.mean(dim=1)
    return embeddings

# ============================ Initialisation du modèle Mistral ============================
try:
    model_chat = ChatMistralAI(model="mistral-large-latest", api_key=key_Mistral)
except Exception as e:
    print(f"Erreur lors de l'initialisation du modèle Mistral: {e}")
    exit()

# ============================ Requête avec gestion des erreurs ============================
def invoke_with_retry(model, prompt, max_retries=10):
    retries = 0
    while retries < max_retries:
        try:
            return model.invoke(prompt)
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 429:
                wait_time = 2 ** retries
                print(f"Rate limit exceeded. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                retries += 1
            else:
                raise e
    raise Exception("Max retries reached, request failed.")

# ============================ Prompt Template ============================
PROMPT_TEMPLATE = """
En tant qu'expert en analyse de finances publiques, rédige une note d'analyse complète de 1500 mots, se concentrant sur les prévisions et hypothèses entourant les hypothèses de 
déficit public et de dette publique, basée sur le contexte suivant :
{context}

Voici un exemple de note d'analyse pour l'année {annee} :
{exemple}

Instructions spécifiques :
1. Structurez votre analyse en sections claires (Contexte économique, Hypothèses techniques, Déficit public, Dépenses publiques, Recettes publiques, Dette publique, Perspectives)
2. Mettez en évidence les points clés et les chiffres importants
3. Utilisez un style professionnel et concis
4. Incluez des tableaux de données pertinents
5. N'utiliser que les informations et chiffres contenus dans le contexte
6. N'utilisez pas les mots "hypothèses techniques" ou "hypothèses de recettes" dans votre analyse
Contexte fourni :
{context}

Rédigez votre analyse en français, en vous basant uniquement sur les informations fournies dans le contexte, mais en vous inspirant de l'exemple fourni pour la structure et le style.
"""
prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)

# ============================ Traitement des documents ============================
path = "C:/Users/menar/Documents/projet_certifIA_banque"
input_dir = path + "/output/textes_extraits"
input_exemple = path + "/output/notes_RESF/textes_extraits"
output_dir = path + "/output/notes_RESF_generees"

os.makedirs(output_dir, exist_ok=True)

# Liste des années à analyser
annees = ["2020", "2021", "2022"]

for annee in annees:
    input_file = f"RESF_{annee}.txt"
    input_path = os.path.join(input_dir, input_file)
    
    # Vérifier si le fichier existe
    if not os.path.exists(input_path):
        print(f"⚠️ Fichier {input_file} non trouvé")
        continue

    # Lire le contenu du fichier
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    document = LangDocument(page_content=content, metadata={"source": input_file})
    
    # Découpage du texte en chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = text_splitter.split_documents([document]) 
    
    # Création de la base de données vectorielle
    chroma_path = f"{path}/output/database_RAG_analyse/db_RESF_{annee}"
    embeddings = MistralAIEmbeddings(model="mistral-embed", api_key=key_Mistral)

    if not os.path.exists(chroma_path):
        db_chroma = Chroma.from_documents(chunks, embeddings, persist_directory=chroma_path)
    else:
        db_chroma = Chroma(persist_directory=chroma_path, embedding_function=embeddings)
    
    # Recherche des passages pertinents
    question = "Analysez la situation financière et économique de la France pour l'année " + annee
    docs_chroma = db_chroma.similarity_search_with_score(question, k=10)
    context_text = "\n\n".join([doc.page_content for doc, _score in docs_chroma])

    # === Ajout : lecture de la note d'exemple ===
    exemple_file = f"Analyse_RESF_{annee}.txt"
    exemple_path = os.path.join(input_exemple, exemple_file)
    if os.path.exists(exemple_path):
        with open(exemple_path, "r", encoding="utf-8") as f:
            exemple_text = f.read()
    else:
        exemple_text = ""

    # Génération de l'analyse
    prompt = prompt_template.format(context=context_text, annee=annee, exemple=exemple_text, question=question)
    response_text = invoke_with_retry(model_chat, prompt)

    # Sauvegarde de l'analyse (texte)
    output_path_txt = os.path.join(output_dir, f"Analyse_RESF_{annee}_generee.txt") 
    with open(output_path_txt, "w", encoding="utf-8") as f:
        f.write(response_text.content)

    # Sauvegarde de l'analyse sous format Word
    output_path_word = os.path.join(output_dir, f"Analyse_RESF_{annee}_generee.docx")
    doc = WordDocument()
    
    # Définition des styles pour les titres
    styles = doc.styles
    style_titre1 = styles['Heading 1']
    style_titre1.font.size = Pt(16)
    style_titre1.font.bold = True
    style_titre1.font.color.rgb = RGBColor(0, 0, 0)
    
    style_titre2 = styles['Heading 2']
    style_titre2.font.size = Pt(14)
    style_titre2.font.bold = True
    style_titre2.font.color.rgb = RGBColor(0, 0, 0)
    
    style_titre3 = styles['Heading 3']
    style_titre3.font.size = Pt(12)
    style_titre3.font.bold = True
    style_titre3.font.color.rgb = RGBColor(0, 0, 0)

    # Titre principal (centré)
    title = doc.add_heading(f'Analyse RESF {annee}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Traitement du texte ligne par ligne
    lines = response_text.content.split('\n')
    current_paragraph = None
    table_lines = []
    in_table = False
    current_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            # Si on a du texte en attente, on le met dans un paragraphe
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            continue
        if line.startswith('#### '):
            # On écrit le texte en attente avant de changer de section
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            # Titre niveau 4
            doc.add_heading(line[4:], level=4)

        elif line.startswith('### '):
            # On écrit le texte en attente avant de changer de section
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            # Titre niveau 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            # On écrit le texte en attente avant de changer de section
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            # Titre niveau 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            # On écrit le texte en attente avant de changer de section
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            # Titre niveau 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('|'):
            # On écrit le texte en attente avant de créer le tableau
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            # C'est un tableau
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            # Si on était dans un tableau, on le crée
            if in_table and table_lines:
                # Création du tableau
                rows = len(table_lines)
                cols = len(table_lines[0].split('|')) - 2  # -2 car les | en début et fin
                
                # Vérification de la taille du tableau
                if rows > 0 and cols > 0:
                    try:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        # Remplissage du tableau
                        for i, table_line in enumerate(table_lines):
                            cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                            for j, cell in enumerate(cells):
                                if j < cols:  # Vérification supplémentaire
                                    # Style de l'en-tête
                                    if i == 0:
                                        paragraph = table.cell(i, j).paragraphs[0]
                                        run = paragraph.add_run(cell)
                                        run.bold = True
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        table.cell(i, j).text = cell
                        
                        # Ajout d'un espace après le tableau
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"⚠️ Erreur lors de la création du tableau : {e}")
                        # On continue avec le texte normal
                in_table = False
                table_lines = []
            
            # Texte normal
            current_text.append(line)

    # Gestion du dernier tableau s'il existe
    if in_table and table_lines:
        rows = len(table_lines)
        cols = len(table_lines[0].split('|')) - 2
        
        # Vérification de la taille du tableau
        if rows > 0 and cols > 0:
            try:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for i, table_line in enumerate(table_lines):
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    for j, cell in enumerate(cells):
                        if j < cols:  # Vérification supplémentaire
                            if i == 0:
                                paragraph = table.cell(i, j).paragraphs[0]
                                run = paragraph.add_run(cell)
                                run.bold = True
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                table.cell(i, j).text = cell
            except Exception as e:
                print(f"⚠️ Erreur lors de la création du dernier tableau : {e}")
                # On continue avec le texte normal

    # Gestion du texte restant
    if current_text:
        if not current_paragraph:
            current_paragraph = doc.add_paragraph()
            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        current_paragraph.add_run(' '.join(current_text))

    doc.save(output_path_word)
    print(f"✅ Analyse générée pour {input_file} → {output_path_word}") 