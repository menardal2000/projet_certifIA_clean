import os
import time
import torch
import httpx
import requests
import pandas as pd
from PIL import Image
from io import BytesIO
from mistralai import Mistral
import datauri  
from mistralai.models import OCRResponse
from transformers import AutoTokenizer, AutoModel
from huggingface_hub import login
from langchain.schema import Document as LangDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.prompts import ChatPromptTemplate
from langchain_mistralai import MistralAIEmbeddings, ChatMistralAI
from docx import Document as WordDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from credentials.key import key_Mistral, HF_TOKEN
from parametres.params import parent_dir, input_dir, output_dir, exemple_dir, notes_generees_dir, database_dir, RESF_dir, docs_dir, tables_dir, txt_dir

# ============================ Initialisation des clés API ============================
os.environ["HF_TOKEN"] = HF_TOKEN
login(os.environ["HF_TOKEN"])

# ============================ Initialisation des clients ============================
client = Mistral(api_key=key_Mistral)
try:
    model_chat = ChatMistralAI(model="mistral-large-latest", api_key=key_Mistral)
except Exception as e:
    print(f"Erreur lors de l'initialisation du modèle Mistral: {e}")
    exit()

# ============================ Modèle d'embeddings local ============================
model = AutoModel.from_pretrained("BAAI/bge-m3")
tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-m3")

def get_embeddings(texts):
    inputs = tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
    with torch.no_grad():
        embeddings = model(**inputs).last_hidden_state.mean(dim=1)
    return embeddings

# ============================ Fonctions utilitaires ============================
def invoke_with_retry(model, prompt, max_retries=10):
    retries = 0
    while retries < max_retries:
        try:
            return model.invoke(prompt)
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 429:
                wait_time = 2 ** retries
                print(f"Rate limit exceeded. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
                retries += 1
            else:
                raise e
    raise Exception("Max retries reached, request failed.")

# ============================ Prompt Template ============================
PROMPT_TEMPLATE = """
En tant qu'expert en analyse de finances publiques, rédige une note d'analyse complète de 1500 mots, se concentrant sur les prévisions et hypothèses entourant les prévisions de 
déficit public et de dette publique, basée sur le contexte suivant :
{context}

Voici un exemple de note d'analyse pour l'année {annee} :
{exemple}

Instructions spécifiques :
1. Structurez votre analyse en sections claires (Contexte économique, Hypothèses techniques, Déficit public, Dépenses publiques, Recettes publiques, Dette publique, Perspectives)
2. Mettez en évidence les points clés et les chiffres importants
3. Utilisez un style professionnel et concis
4. Incluez des tableaux de données pertinents
5. N'utiliser que les informations et chiffres contenus dans le contexte
6. N'utilisez pas les mots "hypothèses techniques" ou "hypothèses de recettes" dans votre analyse
Contexte fourni :
{context}

Rédigez votre analyse en français, en vous basant uniquement sur les informations fournies dans le contexte, mais en vous inspirant de l'exemple fourni pour la structure et le style.
"""
prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)

# ============================ Définition des chemins ============================
path_input = os.path.join(parent_dir, input_dir)
path_resf = os.path.join(path_input, RESF_dir)
path_output = os.path.join(parent_dir, output_dir)
path_exemple = os.path.join(path_input, exemple_dir)
path_notes_generees = os.path.join(path_output, notes_generees_dir)
path_notes_generees_docs = os.path.join(path_notes_generees, docs_dir)
path_notes_generees_tables = os.path.join(path_notes_generees, tables_dir)
path_notes_generees_txt = os.path.join(path_notes_generees, txt_dir)
path_database = os.path.join(path_output, database_dir)

# Création des répertoires
os.makedirs(path_input, exist_ok=True)
os.makedirs(path_output, exist_ok=True)
os.makedirs(path_exemple, exist_ok=True)
os.makedirs(path_notes_generees, exist_ok=True)
os.makedirs(path_notes_generees_docs, exist_ok=True)
os.makedirs(path_notes_generees_tables, exist_ok=True)
os.makedirs(path_database, exist_ok=True)
os.makedirs(path_notes_generees_txt, exist_ok=True)
os.makedirs(path_resf, exist_ok=True)

# ============================ Fonction de traitement OCR ============================
def process_ocr(annee):
    input_file = f"RESF_{annee}.pdf"
    input_path = os.path.join(path_resf, input_file)
    
    if not os.path.exists(input_path):
        print(f"⚠️ Fichier {input_file} non trouvé dans {input_path}")
        return None

    print(f"📄 Traitement du fichier {input_file}...")
    
    try:
        uploaded_pdf = client.files.upload(
            file={
                "file_name": input_path,
                "content": open(input_path, "rb"),
            },
            purpose="ocr"
        )

        signed_url = client.files.get_signed_url(file_id=uploaded_pdf.id)

        ocr_response = client.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "document_url",
                "document_url": signed_url.url,
            }
        )

        ocr_text_with_pages = ""
        for page_number, page in enumerate(ocr_response.pages, start=1):
            if page.markdown:
                ocr_text_with_pages += f"[{page_number}]\n{page.markdown}\n"

        file_name = f"RESF_{annee}.txt"
        file_path = os.path.join(path_notes_generees_txt, file_name)
        with open(file_path, "w", encoding="utf-8") as file:
            file.write(ocr_text_with_pages)
        
        print(f"✅ OCR terminé pour {input_file}")
        return file_path
        
    except Exception as e:
        print(f"❌ Erreur lors du traitement OCR de {input_file}: {str(e)}")
        return None

def extract_and_save_tables(table_lines, annee, table_index):
    """Extrait les tableaux du texte et les sauvegarde en format Excel."""
    if not table_lines:
        return None
        
    # Convertir les lignes du tableau en DataFrame
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    
    if len(rows) < 2:  # Au moins l'en-tête et une ligne de données
        return None
    
    # Vérifier la cohérence du nombre de colonnes
    max_cols = max(len(row) for row in rows)
    min_cols = min(len(row) for row in rows)
    
    if max_cols != min_cols:
        print(f"⚠️ Attention : Le tableau a un nombre de colonnes incohérent (min: {min_cols}, max: {max_cols})")
        # Ajuster toutes les lignes pour avoir le même nombre de colonnes
        for i in range(len(rows)):
            if len(rows[i]) < max_cols:
                rows[i].extend([''] * (max_cols - len(rows[i])))
    
    try:
        # Créer le DataFrame
        df = pd.DataFrame(rows[1:], columns=rows[0])
        
        # Sauvegarder en Excel
        excel_path = os.path.join(path_notes_generees_tables, f"Tableau_RESF_{annee}_{table_index}.xlsx")
        df.to_excel(excel_path, index=False)
        print(f"✅ Tableau sauvegardé avec {len(df.columns)} colonnes : {excel_path}")
        return excel_path
    except Exception as e:
        print(f"❌ Erreur lors de la sauvegarde du tableau : {str(e)}")
        return None

# ============================ Fonction de génération d'analyse ============================
def generate_analysis(annee, input_path):
    # Lecture du contenu du fichier
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    document = LangDocument(page_content=content, metadata={"source": input_path})
    
    # Découpage du texte en chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
    chunks = text_splitter.split_documents([document]) 
    
    # Création de la base de données vectorielle
    chroma_path = os.path.join(path_database, f"db_RESF_{annee}")
    embeddings = MistralAIEmbeddings(model="mistral-embed", api_key=key_Mistral)

    if not os.path.exists(chroma_path):
        db_chroma = Chroma.from_documents(chunks, embeddings, persist_directory=chroma_path)
    else:
        db_chroma = Chroma(persist_directory=chroma_path, embedding_function=embeddings)
    
    # Recherche des passages pertinents
    question = f"Analysez la situation économique et des finances publiques de la France pour l'année {annee}"
    docs_chroma = db_chroma.similarity_search_with_score(question, k=10)
    context_text = "\n\n".join([doc.page_content for doc, _score in docs_chroma])

    # Lecture de la note d'exemple
    exemple_file = f"Analyse_RESF_{annee}.txt"
    exemple_path = os.path.join(path_exemple, exemple_file)
    if os.path.exists(exemple_path):
        with open(exemple_path, "r", encoding="utf-8") as f:
            exemple_text = f.read()
    else:
        exemple_text = ""

    # Génération de l'analyse
    prompt = prompt_template.format(context=context_text, annee=annee, exemple=exemple_text, question=question)
    response_text = invoke_with_retry(model_chat, prompt)

    # Sauvegarde de l'analyse (texte)
    output_path_txt = os.path.join(path_notes_generees_txt, f"Analyse_RESF_{annee}_generee.txt") 
    with open(output_path_txt, "w", encoding="utf-8") as f:
        f.write(response_text.content)

    # Sauvegarde de l'analyse sous format Word
    output_path_word = os.path.join(path_notes_generees_docs, f"Analyse_RESF_{annee}_generee.docx")
    doc = WordDocument()
    
    # Définition des styles pour les titres
    styles = doc.styles
    style_titre1 = styles['Heading 1']
    style_titre1.font.size = Pt(16)
    style_titre1.font.bold = True
    style_titre1.font.color.rgb = RGBColor(0, 0, 0)
    
    style_titre2 = styles['Heading 2']
    style_titre2.font.size = Pt(14)
    style_titre2.font.bold = True
    style_titre2.font.color.rgb = RGBColor(0, 0, 0)
    
    style_titre3 = styles['Heading 3']
    style_titre3.font.size = Pt(12)
    style_titre3.font.bold = True
    style_titre3.font.color.rgb = RGBColor(0, 0, 0)

    # Titre principal (centré)
    title = doc.add_heading(f'Analyse RESF {annee}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Traitement du texte ligne par ligne
    lines = response_text.content.split('\n')
    current_paragraph = None
    table_lines = []
    in_table = False
    current_text = []
    table_index = 1
    
    for line in lines:
        line = line.strip()
        if not line:
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            continue

        if line.startswith('#### '):
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            doc.add_heading(line[4:], level=4)
        elif line.startswith('### '):
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            doc.add_heading(line[2:], level=1)
        elif line.startswith('|'):
            if current_text:
                if not current_paragraph:
                    current_paragraph = doc.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                current_paragraph.add_run(' '.join(current_text))
                current_text = []
                current_paragraph = None
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
        else:
            if in_table and table_lines:
                # Sauvegarder le tableau en Excel
                excel_path = extract_and_save_tables(table_lines, annee, table_index)
                if excel_path:
                    print(f"✅ Tableau sauvegardé : {excel_path}")
                    table_index += 1
                
                # Créer le tableau dans le document Word
                rows = len(table_lines)
                cols = len(table_lines[0].split('|')) - 2
                
                if rows > 0 and cols > 0:
                    try:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.style = 'Table Grid'
                        
                        for i, table_line in enumerate(table_lines):
                            cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                            for j, cell in enumerate(cells):
                                if j < cols:
                                    if i == 0:
                                        paragraph = table.cell(i, j).paragraphs[0]
                                        run = paragraph.add_run(cell)
                                        run.bold = True
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        table.cell(i, j).text = cell
                        
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"⚠️ Erreur lors de la création du tableau : {e}")
                in_table = False
                table_lines = []
            
            current_text.append(line)

    # Gestion du dernier tableau
    if in_table and table_lines:
        # Sauvegarder le dernier tableau en Excel
        excel_path = extract_and_save_tables(table_lines, annee, table_index)
        if excel_path:
            print(f"✅ Dernier tableau sauvegardé : {excel_path}")
        
        # Créer le tableau dans le document Word
        rows = len(table_lines)
        cols = len(table_lines[0].split('|')) - 2
        
        if rows > 0 and cols > 0:
            try:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                for i, table_line in enumerate(table_lines):
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    for j, cell in enumerate(cells):
                        if j < cols:
                            if i == 0:
                                paragraph = table.cell(i, j).paragraphs[0]
                                run = paragraph.add_run(cell)
                                run.bold = True
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                table.cell(i, j).text = cell
            except Exception as e:
                print(f"⚠️ Erreur lors de la création du dernier tableau : {e}")

    # Gestion du texte restant
    if current_text:
        if not current_paragraph:
            current_paragraph = doc.add_paragraph()
            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        current_paragraph.add_run(' '.join(current_text))

    doc.save(output_path_word)
    print(f"✅ Analyse générée pour RESF_{annee} → {output_path_word}")

# ============================ Programme principal ============================
def main():
    annees = ["2020", "2021", "2022", "2023", "2024", "2025"]
    
    for annee in annees:
        print(f"\nTraitement de l'année {annee}...")
        
        # Étape 1: OCR
        input_path = process_ocr(annee)
        if input_path is None:
            continue
            
        # Étape 2: Génération de l'analyse
        generate_analysis(annee, input_path)

if __name__ == "__main__":
    main()